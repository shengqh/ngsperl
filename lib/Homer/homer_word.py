import os
from os.path import basename, dirname
from docx import Document
from docx.shared import Inches
from cairosvg import svg2png
import argparse

# Parse command line arguments
parser = argparse.ArgumentParser(description='Create Word document with Homer motif results')
parser.add_argument('--results_dir', required=True, help='Directory containing the homer results')
parser.add_argument('--target_dir', required=True, help='Directory to save output files')
args = parser.parse_args()

# Directory containing the homer results
results_dir = args.results_dir
target_dir = args.target_dir

# Create target directory if it doesn't exist
os.makedirs(target_dir, exist_ok=True)

doc_name=basename(dirname(results_dir)).split("_vs_")[0]  # Extract the base name for the document

doc = Document()
heading = doc.add_heading(f"TFs enriched in\n{doc_name} ATAC peaks", 3)
heading.alignment = 1  # 1 = center alignment

# Process motifs 1 to 10
for i in range(1, 11):
  motif_file = os.path.join(results_dir, f"known{i}.motif")
  svg_file = os.path.join(results_dir, f"known{i}.logo.svg")
  png_file = os.path.join(target_dir, f"{doc_name}_known{i}.logo.png")
  
  # Check if motif file exists
  if not os.path.exists(motif_file):
    print(f"Motif file {motif_file} does not exist. Skipping motif {i}.")
    continue
  
  if not os.path.exists(svg_file):
    print(f"SVG file {svg_file} does not exist. Skipping motif {i}.")
    continue
  
  # Extract motif name from .motif file (assuming it's on the first line after '>')
  motif_name = ""
  with open(motif_file, 'r') as f:
    motif_name = f.readline().strip().split('\t')[1]  # Get the first line and split by tab
    motif_name = motif_name.split('/')[0]
  
  # Convert SVG to PNG using inkscape (if available)
  if os.path.exists(svg_file):
    try:
      svg2png(url=svg_file, write_to=png_file)
    except:
      print(f"Could not convert {svg_file} to PNG. Make sure inkscape is installed.")
      continue
    
  # Add motif information to Word document
  if os.path.exists(png_file):
    # Create or get the table (create on first iteration)
    if i == 1:
      table = doc.add_table(rows=1, cols=3)
      table.style = 'Table Grid'
      # Add header row
      hdr_cells = table.rows[0].cells
      hdr_cells[0].text = 'Rank'
      hdr_cells[1].text = 'Motif'
      hdr_cells[2].text = 'Name'
      
      # Set specific column widths
      table.columns[0].width = Inches(0.5)  # Rank column
      table.columns[1].width = Inches(4.0)  # Motif column
      table.columns[2].width = Inches(2)  # Name column
      
    # Add a new row to the table
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[2].text = motif_name
    
    # Add image to column 2
    paragraph = row_cells[1].paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.add_picture(png_file, width=Inches(2))

# Save the Word document
doc_file = os.path.join(target_dir, doc_name + "_homer_motifs.docx")
doc.save(doc_file)
print(f"Word document created: {doc_file}")